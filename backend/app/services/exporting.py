from __future__ import annotations

from collections import OrderedDict
import json
from pathlib import Path
import re
import uuid

from docx import Document
from docx.shared import Inches
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.config import Settings
from app.models import CitationSlot, DraftSection, EvidenceMatch, ExportBundle, FigureAsset, FigureSpec, Outline, Project, QualityReport, ReferenceRecord


CITATION_GROUP_RE = re.compile(r"\[([^\]]*CIT_[A-Z0-9_][^\]]*)\]")
CITATION_TOKEN_RE = re.compile(r"\b(CIT_[A-Z0-9_]+)\b")
FIGURE_PLACEHOLDER_RE = re.compile(r"\[FIGURE[_ ](\d+):\s*([^\]]+?)\]")


def _latest_sections(session: Session, project_id: str) -> list[DraftSection]:
    sections = list(session.scalars(select(DraftSection).where(DraftSection.project_id == project_id).order_by(DraftSection.created_at)).all())
    latest: dict[str, DraftSection] = {}
    for section in sections:
        current = latest.get(section.section_key)
        if current is None or section.version > current.version:
            latest[section.section_key] = section
    return list(latest.values())


def _latest_outline(session: Session, project_id: str) -> Outline | None:
    return session.scalars(select(Outline).where(Outline.project_id == project_id).order_by(Outline.version.desc())).first()


def _latest_quality_report(session: Session, project_id: str) -> QualityReport | None:
    return session.scalars(
        select(QualityReport).where(QualityReport.project_id == project_id).order_by(QualityReport.version.desc())
    ).first()


def _extract_slot_keys(content: str) -> list[str]:
    keys: list[str] = []
    seen: set[str] = set()
    for slot_key in CITATION_TOKEN_RE.findall(content):
        if slot_key in seen:
            continue
        seen.add(slot_key)
        keys.append(slot_key)
    return keys


def _citation_numbering(session: Session, project_id: str) -> tuple[dict[str, int], dict[int, ReferenceRecord]]:
    slots = {
        slot.slot_key: slot
        for slot in session.scalars(select(CitationSlot).where(CitationSlot.project_id == project_id)).all()
    }
    matches = {
        match.citation_slot_id: match
        for match in session.scalars(select(EvidenceMatch).where(EvidenceMatch.project_id == project_id)).all()
    }
    sections = _latest_sections(session, project_id)
    ordered_refs: OrderedDict[str, ReferenceRecord] = OrderedDict()
    for section in sections:
        for slot_key in _extract_slot_keys(section.content):
            slot = slots.get(slot_key)
            if slot is None:
                continue
            match = matches.get(slot.id)
            if match is None:
                continue
            for ref_id in match.selected_reference_ids_json:
                if ref_id not in ordered_refs:
                    reference = session.get(ReferenceRecord, ref_id)
                    if reference is not None:
                        ordered_refs[ref_id] = reference
    slot_numbers: dict[str, int] = {}
    numbered_refs: dict[int, ReferenceRecord] = {}
    for index, (ref_id, reference) in enumerate(ordered_refs.items(), start=1):
        numbered_refs[index] = reference
        for slot in slots.values():
            match = matches.get(slot.id)
            if match and ref_id in match.selected_reference_ids_json and slot.slot_key not in slot_numbers:
                slot_numbers[slot.slot_key] = index
    return slot_numbers, numbered_refs


def _strip_duplicate_heading(content: str, heading: str) -> str:
    lines = content.splitlines()
    if not lines:
        return content.strip()
    normalized_heading = heading.strip().lower()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines:
        candidate = lines[0].strip()
        if candidate.startswith("#"):
            candidate = candidate.lstrip("#").strip()
        if candidate.lower() == normalized_heading:
            lines = lines[1:]
    return "\n".join(lines).strip()


def _render_citation_group(group_text: str, slot_numbers: dict[str, int]) -> str:
    rendered_parts: list[str] = []
    seen: set[str] = set()
    for slot_key in _extract_slot_keys(group_text):
        number = slot_numbers.get(slot_key)
        label = "manual review" if number is None else str(number)
        if label in seen:
            continue
        seen.add(label)
        rendered_parts.append(label)
    if not rendered_parts:
        return "[manual review]"
    return f"[{', '.join(rendered_parts)}]"


def _render_content(content: str, heading: str, slot_numbers: dict[str, int]) -> str:
    rendered = _strip_duplicate_heading(content, heading)

    rendered = CITATION_GROUP_RE.sub(lambda match: _render_citation_group(match.group(1), slot_numbers), rendered)
    rendered = CITATION_TOKEN_RE.sub(
        lambda match: f"[{slot_numbers[match.group(1)]}]" if match.group(1) in slot_numbers else "[manual review]",
        rendered,
    )
    rendered = re.sub(r"`(\[[^`]+\])`", r"\1", rendered)
    rendered = FIGURE_PLACEHOLDER_RE.sub(
        lambda match: f"\nFigure {match.group(1)}. Suggested insert: {match.group(2).strip()}\n",
        rendered,
    )
    rendered = re.sub(r"\n{3,}", "\n\n", rendered)
    return rendered.strip()


def _content_paragraphs(content: str) -> list[str]:
    return [paragraph.strip() for paragraph in re.split(r"\n\s*\n", content) if paragraph.strip()]


def _figure_details(session: Session, project_id: str) -> dict[int, dict[str, object]]:
    specs = list(
        session.scalars(select(FigureSpec).where(FigureSpec.project_id == project_id).order_by(FigureSpec.figure_number)).all()
    )
    details: dict[int, dict[str, object]] = {}
    for spec in specs:
        asset = next((item for item in spec.figure_assets if item.selected and item.artifact is not None), None)
        details[spec.figure_number] = {
            "figure_key": spec.figure_key,
            "caption": spec.caption_draft,
            "artifact_path": "" if asset is None else asset.artifact.storage_path,
            "asset_id": None if asset is None else asset.id,
            "method_section_content": spec.method_section_content,
        }
    return details


def _render_text_only(paragraph: str, slot_numbers: dict[str, int]) -> str:
    rendered = CITATION_GROUP_RE.sub(lambda match: _render_citation_group(match.group(1), slot_numbers), paragraph)
    rendered = CITATION_TOKEN_RE.sub(
        lambda match: f"[{slot_numbers[match.group(1)]}]" if match.group(1) in slot_numbers else "[manual review]",
        rendered,
    )
    rendered = re.sub(r"`(\[[^`]+\])`", r"\1", rendered)
    rendered = re.sub(r"\n{3,}", "\n\n", rendered)
    return rendered.strip()


def _render_blocks(
    content: str,
    heading: str,
    slot_numbers: dict[str, int],
    figure_details: dict[int, dict[str, object]],
) -> list[dict[str, object]]:
    normalized = _strip_duplicate_heading(content, heading)
    blocks: list[dict[str, object]] = []
    for paragraph in _content_paragraphs(normalized):
        match = FIGURE_PLACEHOLDER_RE.fullmatch(paragraph.strip())
        if match:
            figure_number = int(match.group(1))
            figure_info = figure_details.get(figure_number)
            blocks.append(
                {
                    "type": "figure",
                    "figure_number": figure_number,
                    "caption": match.group(2).strip() if figure_info is None else str(figure_info["caption"]),
                    "artifact_path": None if figure_info is None else figure_info["artifact_path"],
                }
            )
        else:
            blocks.append({"type": "text", "content": _render_text_only(paragraph, slot_numbers)})
    return blocks


def _render_reference_line(index: int, reference: ReferenceRecord) -> str:
    authors = ", ".join(reference.authors_json[:3]) if reference.authors_json else "Unknown author"
    year = f" ({reference.year})" if reference.year else ""
    venue = f". {reference.venue}" if reference.venue else ""
    doi = f". doi:{reference.doi}" if reference.doi else ""
    return f"[{index}] {authors}{year}. {reference.title}{venue}{doi}".strip()


def _render_bibtex_entry(reference: ReferenceRecord) -> str:
    key = (reference.doi or reference.external_id or reference.id).replace("/", "_").replace(":", "_")
    authors = " and ".join(reference.authors_json) if reference.authors_json else "Unknown"
    return (
        f"@article{{{key},\n"
        f"  title = {{{reference.title}}},\n"
        f"  author = {{{authors}}},\n"
        f"  journal = {{{reference.venue}}},\n"
        f"  year = {{{reference.year or ''}}},\n"
        f"  doi = {{{reference.doi}}},\n"
        f"}}"
    )


def run_export(session: Session, project: Project, settings: Settings, *, mode: str = "draft") -> ExportBundle:
    if mode == "final":
        report = _latest_quality_report(session, project.id)
        if report is None or not report.submission_ready:
            raise ValueError("Final export is blocked until all critical quality issues are resolved")

    output_root = Path(settings.storage_root).expanduser().resolve() / "exports" / project.id / str(uuid.uuid4())
    output_root.mkdir(parents=True, exist_ok=True)

    slot_numbers, numbered_refs = _citation_numbering(session, project.id)
    outline = _latest_outline(session, project.id)
    sections = _latest_sections(session, project.id)
    figure_details = _figure_details(session, project.id)

    md_lines = [f"# {project.title}", ""]
    if outline is not None:
        md_lines.append(f"_Type: {outline.manuscript_type}_")
        md_lines.append("")
    for section in sections:
        md_lines.append(f"## {section.heading}")
        md_lines.append("")
        for block in _render_blocks(section.content, section.heading, slot_numbers, figure_details):
            if block["type"] == "text":
                md_lines.append(str(block["content"]))
                md_lines.append("")
                continue
            figure_number = int(block["figure_number"])
            caption = str(block["caption"])
            artifact_path = str(block["artifact_path"] or "")
            if artifact_path:
                md_lines.append(f"![Figure {figure_number}]({artifact_path})")
                md_lines.append("")
                md_lines.append(f"Figure {figure_number}. {caption}")
            else:
                md_lines.append(f"Figure {figure_number}. Suggested insert: {caption}")
            md_lines.append("")
    md_lines.append("## References")
    md_lines.append("")
    for index, reference in numbered_refs.items():
        md_lines.append(_render_reference_line(index, reference))
    markdown = "\n".join(md_lines).strip() + "\n"

    bibtex = "\n\n".join(_render_bibtex_entry(reference) for reference in numbered_refs.values()).strip() + "\n"
    manuscript_json = {
        "project": {"id": project.id, "title": project.title, "objective": project.objective},
        "sections": [
            {
                "section_key": section.section_key,
                "heading": section.heading,
                "content": "\n\n".join(
                    str(block["content"])
                    if block["type"] == "text"
                    else (
                        f"Figure {int(block['figure_number'])}. {str(block['caption'])}"
                        if block["artifact_path"]
                        else f"Figure {int(block['figure_number'])}. Suggested insert: {str(block['caption'])}"
                    )
                    for block in _render_blocks(section.content, section.heading, slot_numbers, figure_details)
                ),
            }
            for section in sections
        ],
        "references": [
            {
                "number": index,
                "title": reference.title,
                "doi": reference.doi,
                "authors": reference.authors_json,
            }
            for index, reference in numbered_refs.items()
        ],
        "figures": [
            {
                "figure_number": figure_number,
                "artifact_path": details["artifact_path"],
                "caption": details["caption"],
                "method_section_content": details["method_section_content"],
            }
            for figure_number, details in sorted(figure_details.items())
        ],
    }

    md_path = output_root / "manuscript.md"
    md_path.write_text(markdown, encoding="utf-8")
    bib_path = output_root / "references.bib"
    bib_path.write_text(bibtex, encoding="utf-8")
    json_path = output_root / "manuscript.json"
    json_path.write_text(json.dumps(manuscript_json, indent=2), encoding="utf-8")

    document = Document()
    document.add_heading(project.title, level=0)
    for section in sections:
        document.add_heading(section.heading, level=1)
        for block in _render_blocks(section.content, section.heading, slot_numbers, figure_details):
            if block["type"] == "text":
                document.add_paragraph(str(block["content"]))
                continue
            artifact_path = str(block["artifact_path"] or "")
            if artifact_path and Path(artifact_path).exists():
                document.add_picture(artifact_path, width=Inches(6.5))
                document.add_paragraph(f"Figure {int(block['figure_number'])}. {str(block['caption'])}")
            else:
                document.add_paragraph(f"Figure {int(block['figure_number'])}. Suggested insert: {str(block['caption'])}")
    document.add_heading("References", level=1)
    for index, reference in numbered_refs.items():
        document.add_paragraph(_render_reference_line(index, reference))
    docx_path = output_root / "manuscript.docx"
    document.save(docx_path)

    bundle = ExportBundle(
        project_id=project.id,
        status="ready",
        manifest_json={
            "mode": mode,
            "markdown_path": str(md_path),
            "bibtex_path": str(bib_path),
            "json_path": str(json_path),
            "docx_path": str(docx_path),
        },
    )
    session.add(bundle)
    session.commit()
    session.refresh(bundle)
    return bundle
